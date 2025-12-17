from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

output_folder = "reference_files"
os.makedirs(output_folder, exist_ok=True)
docx_filename = "reference.docx"
docx_path = os.path.join(output_folder, docx_filename)


document = Document()

# Set moderate margins
sections = document.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

# --- Define and apply styles programmatically ---
styles = document.styles

# Normal style (for general text)
if 'Normal' not in styles:
    styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.size = Pt(12)
normal_paragraph_format = normal_style.paragraph_format
normal_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
normal_paragraph_format.space_before = Pt(0)
normal_paragraph_format.space_after = Pt(0)
normal_paragraph_format.line_spacing = Pt(14) # Approx 1.15 lines for 12pt font

# Heading 1 style (for book title and TOC title)
if 'Heading 1' not in styles:
    styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
h1_style = styles['Heading 1']
h1_font = h1_style.font
h1_font.size = Pt(16)
h1_font.bold = True
h1_paragraph_format = h1_style.paragraph_format
h1_paragraph_format.space_before = Pt(12) # Add space before
h1_paragraph_format.space_after = Pt(0)
h1_paragraph_format.line_spacing = Pt(18) # Approx 1.15 lines for 16pt font

# Heading 2 style (for chapter titles)
if 'Heading 2' not in styles:
    styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
h2_style = styles['Heading 2']
h2_font = h2_style.font
h2_font.size = Pt(14)
h2_font.bold = True
h2_paragraph_format = h2_style.paragraph_format
h2_paragraph_format.space_before = Pt(12) # Add space before
h2_paragraph_format.space_after = Pt(0)
h2_paragraph_format.line_spacing = Pt(16) # Approx 1.15 lines for 14pt font

# Heading 3 style (for sub-topics)
if 'Heading 3' not in styles:
    styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
h3_style = styles['Heading 3']
h3_font = h3_style.font
h3_font.size = Pt(12)
h3_font.bold = False # Sub-topics are not bold by default, but can be if markdown uses **
h3_paragraph_format = h3_style.paragraph_format
h3_paragraph_format.space_before = Pt(6) # Add smaller space before for sub-topics
h3_paragraph_format.space_after = Pt(0)
h3_paragraph_format.line_spacing = Pt(14) # Approx 1.15 lines for 12pt font

# List Bullet style
if 'List Bullet' not in styles:
    styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
list_bullet_style = styles['List Bullet']
list_bullet_font = list_bullet_style.font
list_bullet_font.size = Pt(12)
list_bullet_paragraph_format = list_bullet_style.paragraph_format
list_bullet_paragraph_format.space_before = Pt(0)
list_bullet_paragraph_format.space_after = Pt(0)
list_bullet_paragraph_format.line_spacing = Pt(14)

# Title Page style (for the main book title)
if 'Title Page' not in styles:
    styles.add_style('Title Page', WD_STYLE_TYPE.PARAGRAPH)
title_page_style = styles['Title Page']
title_page_font = title_page_style.font
title_page_font.size = Pt(24)
title_page_font.bold = True
title_page_paragraph_format = title_page_style.paragraph_format
title_page_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_page_paragraph_format.space_before = Pt(0)
title_page_paragraph_format.space_after = Pt(0)

# --- No sample content added. This file is for styles only. ---

# --- Add page numbers to the header ---
# This is a more robust way to add page numbers, but still can be tricky.
# It inserts a PAGE field and NUMPAGES field.
header = document.sections[0].header
paragraph = header.paragraphs[0]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add page number field
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar1)
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE \\* MERGEFORMAT'
run._r.append(instrText)
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar2)

# Add " of "
paragraph.add_run(' of ')

# Add total pages field
run = paragraph.add_run()
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar3)
instrText2 = OxmlElement('w:instrText')
instrText2.set(qn('xml:space'), 'preserve')
instrText2.text = 'NUMPAGES \\* MERGEFORMAT'
run._r.append(instrText2)
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar4)


document.save(docx_path)
print(f"Created reference DOCX with programmatic styles at: {docx_path}")
print("IMPORTANT: Please open this file in Microsoft Word, verify ALL styles (Heading 1, Heading 2, Heading 3, Normal, List Bullet) and page numbers. Ensure font sizes, bolding, justification, and especially **line spacing (0pt before/after, 14pt/16pt/18pt exactly)** are correct. After verifying, **SAVE THE DOCUMENT** and then proceed to run `main.py`.")
