"""Export structured study notes to Word with cover, TOC, and page breaks."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.modules.export.document_formatter import BookCoverMeta

logger = logging.getLogger(__name__)

_TOC_TAB_POSITION = Inches(6.2)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _insert_field(paragraph, field_code: str, *, placeholder: str = "") -> None:
    """Insert a Word field (PAGE, PAGEREF, etc.) into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    if placeholder:
        text_el = OxmlElement("w:t")
        text_el.text = placeholder
        r.append(text_el)
    r.append(fld_end)


class _BookmarkIds:
    def __init__(self) -> None:
        self._n = 0

    def next(self) -> int:
        self._n += 1
        return self._n


def _add_bookmark(paragraph, name: str, bookmark_ids: _BookmarkIds) -> None:
    """Wrap the first run of a heading paragraph with a named bookmark."""
    if not paragraph.runs:
        paragraph.add_run()
    run = paragraph.runs[0]
    bid = bookmark_ids.next()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    run._r.addprevious(start)
    run._r.addnext(end)


def _sanitize_bookmark_name(raw: str) -> str:
    name = re.sub(r"[^A-Za-z0-9_]", "_", raw.strip())[:40] or "item"
    if name[0].isdigit():
        name = f"b_{name}"
    return name


def _enable_update_fields_on_open(doc: Document) -> None:
    """Ask Word to refresh TOC/page fields when the document is opened."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def _add_page_number_footer(doc: Document) -> None:
    """Centered 'Page X of Y' footer on every page."""
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Page ")
        _insert_field(p, "PAGE", placeholder="1")
        p.add_run(" of ")
        _insert_field(p, "NUMPAGES", placeholder="1")


def _add_toc_line(doc: Document, *, level: int, title: str, bookmark: str) -> None:
    """One TOC row: title ........ page (PAGEREF to heading bookmark)."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.tab_stops.add_tab_stop(_TOC_TAB_POSITION, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if level >= 2:
        fmt.left_indent = Inches(0.4)
    label = p.add_run(title[:110])
    if level == 1:
        label.bold = True
    p.add_run("\t")
    _insert_field(p, f"PAGEREF {bookmark} \\h")


def _prepend_toc_block(
    doc: Document,
    *,
    insert_before,
    toc_rows: Sequence[tuple[int, str, str]],
) -> None:
    """Insert TOC page (title + PAGEREF lines + page break) before first chapter."""
    block: List[Any] = []

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Table of Contents")
    tr.bold = True
    tr.font.size = Pt(16)
    block.append(title._element)

    doc.add_paragraph()
    block.append(doc.paragraphs[-1]._element)

    for level, heading, bookmark in toc_rows:
        _add_toc_line(doc, level=level, title=heading, bookmark=bookmark)
        block.append(doc.paragraphs[-1]._element)

    anchor = insert_before
    for el in reversed(block):
        anchor.addprevious(el)
        anchor = el


def _add_inline_runs(paragraph, text: str) -> None:
    """Add paragraph runs with basic **bold** / *italic* markup."""
    if not text:
        return
    parts: List[tuple[str, bool, bool]] = [(text, False, False)]
    for pattern, bold, italic in ((_BOLD_RE, True, False), (_ITALIC_RE, False, True)):
        new_parts: List[tuple[str, bool, bool]] = []
        for chunk, is_b, is_i in parts:
            if is_b or is_i:
                new_parts.append((chunk, is_b, is_i))
                continue
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    new_parts.append((chunk[pos : m.start()], False, False))
                new_parts.append((m.group(1), bold, italic))
                pos = m.end()
            if pos < len(chunk):
                new_parts.append((chunk[pos:], False, False))
        parts = new_parts or parts
    for chunk, is_b, is_i in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = is_b
        run.italic = is_i


def append_markdown_body(doc: Document, text: str) -> None:
    """Render common markdown (headings, bullets, paragraphs) into the document."""
    if not (text or "").strip():
        return
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        lines = [ln.rstrip() for ln in block.split("\n") if ln.strip()]
        if not lines:
            continue
        first = lines[0].strip()
        if first.startswith("### "):
            doc.add_heading(first[4:].strip(), level=3)
            for ln in lines[1:]:
                _append_line(doc, ln)
            continue
        if first.startswith("## "):
            doc.add_heading(first[3:].strip(), level=2)
            for ln in lines[1:]:
                _append_line(doc, ln)
            continue
        if all(ln.lstrip().startswith(("- ", "* ")) for ln in lines):
            for ln in lines:
                item = re.sub(r"^[\-*]\s+", "", ln.lstrip())
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item)
            continue
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(lines))


def _append_line(doc: Document, line: str) -> None:
    s = line.strip()
    if not s:
        return
    if s.startswith(("- ", "* ")):
        p = doc.add_paragraph(style="List Bullet")
        _add_inline_runs(p, s[2:].strip())
        return
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3)
        return
    p = doc.add_paragraph()
    _add_inline_runs(p, s)


def _add_cover_page(doc: Document, cover: BookCoverMeta) -> None:
    title = (cover.title or "Study Notes").strip()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(cover.subtitle or "Study Notes")
    sub_run.font.size = Pt(14)

    doc.add_paragraph()
    rows = [
        ("Book", title),
        ("Source PDF", cover.source_pdf),
        ("Generated", cover.generated_at),
        ("Chapters", str(cover.chapter_count) if cover.chapter_count else ""),
        ("Sections", str(cover.section_count) if cover.section_count else ""),
        ("Notes style", cover.user_instruction),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for label, val in rows:
        if not (val or "").strip():
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def _add_toc_page(
    doc: Document,
    hierarchy: Dict[str, Any],
    *,
    toc_entries: Optional[Sequence[Any]] = None,
) -> None:
    """Deprecated — TOC is built via _prepend_toc_block after content."""
    return


class DocxNotesExporter:
    """Build a formatted Word document from 15e hierarchy + rewritten section bodies."""

    def export(
        self,
        output_path: str | Path,
        *,
        cover: BookCoverMeta,
        hierarchy: Dict[str, Any],
        rewritten: Dict[str, str],
        reference_docx: Optional[str] = None,
    ) -> str:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        bookmark_ids = _BookmarkIds()
        toc_rows: List[tuple[int, str, str]] = []
        first_chapter_el = None

        _add_cover_page(doc, cover)
        _add_page_break(doc)

        for ch in hierarchy.get("chapters") or []:
            ch_heading = str(ch.get("heading") or "").strip()
            if not ch_heading:
                continue
            sec_blocks: List[tuple[str, str, str]] = []
            for sec in ch.get("sections") or []:
                sid = str(sec.get("section_id") or "")
                body = (rewritten.get(sid) or "").strip()
                if not body:
                    continue
                sec_blocks.append((sid, str(sec.get("heading") or sid), body))
            if not sec_blocks:
                continue

            _add_page_break(doc)
            if first_chapter_el is None:
                first_chapter_el = doc.paragraphs[-1]._element

            ch_id = str(ch.get("chapter_id") or ch_heading)
            ch_bm = _sanitize_bookmark_name(f"ch_{ch_id}")
            h1 = doc.add_heading(ch_heading, level=1)
            _add_bookmark(h1, ch_bm, bookmark_ids)
            toc_rows.append((1, ch_heading, ch_bm))

            for sid, sec_heading, body in sec_blocks:
                sec_bm = _sanitize_bookmark_name(f"sec_{sid}")
                h2 = doc.add_heading(sec_heading[:200], level=2)
                _add_bookmark(h2, sec_bm, bookmark_ids)
                toc_rows.append((2, sec_heading[:110], sec_bm))
                append_markdown_body(doc, body)

        if first_chapter_el is not None and toc_rows:
            _prepend_toc_block(doc, insert_before=first_chapter_el, toc_rows=toc_rows)

        _add_page_number_footer(doc)
        _enable_update_fields_on_open(doc)
        doc.save(str(out))
        _refresh_word_fields(str(out))
        logger.info("Saved formatted Word document: %s", out)
        return str(out)


def _refresh_word_fields(docx_path: str) -> None:
    """Update PAGE/PAGEREF fields via Word on Windows so TOC shows real page numbers."""
    try:
        import win32com.client  # type: ignore[import-untyped]
    except ImportError:
        logger.warning(
            "pywin32 not installed — TOC page numbers stay as placeholders until Word "
            "updates fields (Ctrl+A, F9). Install with: pip install pywin32"
        )
        return

    resolved = str(Path(docx_path).resolve())
    word = None
    doc = None
    wd_print_view = 3
    try:
        word = win32com.client.Dispatch("Word.Application")
        try:
            word.Visible = False
        except AttributeError:
            pass
        word.DisplayAlerts = 0
        doc = word.Documents.Open(resolved)
        doc.ActiveWindow.View.Type = wd_print_view
        doc.Repaginate()
        doc.Fields.Update()
        story = doc.StoryRanges(1)
        while story is not None:
            story.Fields.Update()
            story = story.NextStoryRange
        doc.Fields.Update()
        doc.Save()
        logger.info("Updated Word fields (TOC page numbers): %s", docx_path)
    except Exception as exc:
        logger.warning("Could not auto-update Word fields: %s", exc)
    finally:
        if doc is not None:
            doc.Close(SaveChanges=False)
        if word is not None:
            word.Quit()


def parse_section_bodies_from_markdown(md_text: str) -> Dict[str, str]:
    """Map section heading -> body text from all ## sections in the notes body."""
    m = re.search(r"# Table of Contents[\s\S]*?\n# ([^\n]+)", md_text)
    body = md_text[m.start(1) :] if m else md_text

    sections: Dict[str, str] = {}
    current_heading = ""
    buf: List[str] = []
    skip_block = False

    def flush() -> None:
        nonlocal buf, current_heading
        if current_heading and buf:
            text = "\n".join(buf).strip()
            if text:
                sections[current_heading] = text
        buf = []

    for line in body.splitlines():
        if line.strip().startswith("```{=openxml}"):
            skip_block = True
            continue
        if skip_block:
            if line.strip() == "```":
                skip_block = False
            continue
        if line.startswith("# ") and not line.startswith("## "):
            flush()
            current_heading = ""
            continue
        if line.startswith("## ") and not line.startswith("### "):
            flush()
            current_heading = line[3:].strip()
            if re.match(r"^\d+\.\s+", current_heading):
                current_heading = ""
            continue
        if current_heading:
            buf.append(line)
    flush()
    return sections


def _norm_heading(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def rewritten_map_from_section_bodies(
    hierarchy: Dict[str, Any],
    section_bodies: Dict[str, str],
) -> Dict[str, str]:
    """Match section bodies to section_id by normalized heading."""
    by_norm = {_norm_heading(k): v for k, v in section_bodies.items()}
    result: Dict[str, str] = {}
    for ch in hierarchy.get("chapters") or []:
        for sec in ch.get("sections") or []:
            sid = str(sec.get("section_id") or "")
            heading = str(sec.get("heading") or "").strip()
            body = section_bodies.get(heading) or by_norm.get(_norm_heading(heading), "")
            if body:
                result[sid] = body
    return result
