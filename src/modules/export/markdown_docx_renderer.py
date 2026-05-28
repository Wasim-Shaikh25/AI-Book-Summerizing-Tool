"""Faithful Markdown → DOCX renderer (matches .md structure in Word)."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_OPENXML_PAGE_BREAK = re.compile(r'^\s*<w:p><w:r><w:br w:type="page"/></w:r></w:p>\s*$')
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")


def _add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return
    parts: List[tuple[str, bool, bool]] = [(text, False, False)]
    for pattern, bold, italic in ((_BOLD_RE, True, False), (_ITALIC_RE, False, True)):
        new_parts: List[tuple[str, bool, bool]] = []
        for chunk, is_b, is_i in parts:
            if is_b or is_i:
                new_parts.append((chunk, is_b, is_i))
                continue
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    new_parts.append((chunk[pos : m.start()], False, False))
                new_parts.append((m.group(1), bold, italic))
                pos = m.end()
            if pos < len(chunk):
                new_parts.append((chunk[pos:], False, False))
        parts = new_parts or parts
    for chunk, is_b, is_i in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = is_b
        run.italic = is_i


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _clear_document_body(doc: Document) -> None:
    for p in list(doc.paragraphs):
        el = p._element
        el.getparent().remove(el)
    for t in list(doc.tables):
        el = t._element
        el.getparent().remove(el)


def _bullet_style(level: int) -> str:
    if level <= 0:
        return "List Bullet"
    if level == 1:
        return "List Bullet 2"
    return "List Bullet 3"


def _parse_table_row(line: str) -> List[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{2,}", line))


def render_markdown_to_document(doc: Document, md_text: str) -> None:
    """Append markdown content to an existing document."""
    center_mode = False
    in_code_block = False
    table_rows: List[List[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        while table_rows and all(not (c or "").strip() for c in table_rows[0]):
            table_rows.pop(0)
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(table_rows):
            for ci in range(cols):
                val = row[ci] if ci < len(row) else ""
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_inline_runs(p, val)
        table_rows = []

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if in_code_block:
            if line.strip() == "```":
                in_code_block = False
            elif _OPENXML_PAGE_BREAK.match(line):
                flush_table()
                _add_page_break(doc)
            continue

        if line.strip().startswith("```{=openxml}"):
            flush_table()
            in_code_block = True
            continue

        if line.strip().startswith("|"):
            if _is_table_separator(line):
                continue
            table_rows.append(_parse_table_row(line))
            continue
        flush_table()

        if line.startswith("<div"):
            center_mode = "center" in line.lower()
            continue
        if line.startswith("</div"):
            center_mode = False
            continue

        if not line.strip():
            continue

        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            indent, text = bullet.groups()
            level = min(len(indent) // 2, 2)
            p = doc.add_paragraph(style=_bullet_style(level))
            _add_inline_runs(p, text.strip())
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            indent, text = ordered.groups()
            level = min(len(indent) // 2, 2)
            style = "List Number" if level == 0 else f"List Number {level + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text.strip())
            continue

        p = doc.add_paragraph()
        if center_mode:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line.startswith("#"):
                for run in p.runs:
                    run.bold = True
        _add_inline_runs(p, line.strip())

    flush_table()


def export_markdown_file_to_docx(
    md_text: str,
    output_path: str | Path,
    *,
    reference_docx: Optional[str] = None,
) -> str:
    """Convert a full notes .md file to .docx preserving its structure."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    if reference_docx and Path(reference_docx).exists():
        doc = Document(reference_docx)
        _clear_document_body(doc)
    else:
        doc = Document()

    render_markdown_to_document(doc, md_text)
    doc.save(str(out))
    logger.info("Exported markdown to Word: %s", out)
    return str(out)
