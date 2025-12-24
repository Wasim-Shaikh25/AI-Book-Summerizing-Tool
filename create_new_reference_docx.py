from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING # Import WD_LINE_SPACING
import os

# Define the base directory of the project
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REFERENCE_DOCX_PATH = os.path.join(BASE_DIR, "reference_files", "new_reference.docx")

def create_basic_reference_docx(filepath):
    document = Document()

    # Add some basic styles
    # Normal style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0) # Ensure no space after paragraphs
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE # Ensure single line spacing

    # Heading 1 style
    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(24)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 style
    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(18)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3 style
    style = document.styles['Heading 3']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a few paragraphs with different styles to ensure they are saved
    document.add_heading('This is a Heading 1', level=1)
    document.add_paragraph('This is a normal paragraph with some text to establish the default style.')
    document.add_heading('This is a Heading 2', level=2)
    document.add_paragraph('This is another normal paragraph.')
    document.add_heading('This is a Heading 3', level=3)
    document.add_paragraph('This is a paragraph under Heading 3.')

    document.save(filepath)
    print(f"Created basic reference.docx at {filepath}")

if __name__ == "__main__":
    os.makedirs(os.path.dirname(REFERENCE_DOCX_PATH), exist_ok=True)
    create_basic_reference_docx(REFERENCE_DOCX_PATH)
