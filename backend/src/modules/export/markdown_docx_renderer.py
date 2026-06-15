"""Faithful Markdown → DOCX renderer (matches .md structure in Word)."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

from src.modules.export.docx_theme import (
    NumberedListTracker,
    add_callout_paragraph,
    add_cover_header_band,
    add_topic_subheading,
    apply_study_notes_theme,
    resolve_docx_theme,
    callout_label_text,
    finalize_word_document,
    format_content_table,
    format_metadata_table,
    is_callout_label,
    is_named_callout_label,
    refresh_word_fields,
    style_body_paragraph,
    style_bullet_paragraph,
    style_chapter_heading,
    style_cover_subtitle,
    style_cover_title,
    style_numbered_paragraph,
    style_section_heading,
    style_toc_title,
)

logger = logging.getLogger(__name__)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_OPENXML_PAGE_BREAK = re.compile(r'^\s*<w:p><w:r><w:br w:type="page"/></w:r></w:p>\s*$')
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
_NOTES_STYLE_RE = re.compile(r"^\|\s*\*\*Notes style\*\*\s*\|\s*(.+?)\s*\|", re.I)


def _extract_notes_style_instruction(md_text: str) -> str:
    for line in (md_text or "").splitlines():
        m = _NOTES_STYLE_RE.match(line.strip())
        if m:
            return m.group(1).strip()
    return ""


def _normalize_notes_body_for_export(md_text: str) -> str:
    """Re-apply paragraph/bold layout fixes before Word export (idempotent)."""
    from src.modules.generation.markdown_format_normalizer import strict_normalize_markdown

    instruction = _extract_notes_style_instruction(md_text)
    parts = md_text.split("# Table of Contents")
    if len(parts) < 2:
        return strict_normalize_markdown(md_text, user_instruction=instruction)
    head, tail = parts[0], "# Table of Contents" + parts[1]
    # Normalize only the study body (after TOC), preserve cover/TOC block
    body_split = tail.split("\n# ", 1)
    if len(body_split) < 2:
        return strict_normalize_markdown(md_text, user_instruction=instruction)
    toc_block, body = body_split[0], body_split[1]
    normalized_body = strict_normalize_markdown(f"# {body}", user_instruction=instruction)
    return head + toc_block + "\n# " + normalized_body.lstrip("# ")


def _add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return
    parts: List[tuple[str, bool, bool]] = [(text, False, False)]
    for pattern, bold, italic in ((_BOLD_RE, True, False), (_ITALIC_RE, False, True)):
        new_parts: List[tuple[str, bool, bool]] = []
        for chunk, is_b, is_i in parts:
            if is_b or is_i:
                new_parts.append((chunk, is_b, is_i))
                continue
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    new_parts.append((chunk[pos : m.start()], False, False))
                new_parts.append((m.group(1), bold, italic))
                pos = m.end()
            if pos < len(chunk):
                new_parts.append((chunk[pos:], False, False))
        parts = new_parts or parts
    for chunk, is_b, is_i in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = is_b
        run.italic = is_i


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _clear_document_body(doc: Document) -> None:
    for p in list(doc.paragraphs):
        el = p._element
        el.getparent().remove(el)
    for t in list(doc.tables):
        el = t._element
        el.getparent().remove(el)


def _bullet_style(level: int) -> str:
    if level <= 0:
        return "List Bullet"
    if level == 1:
        return "List Bullet 2"
    return "List Bullet 3"


def _parse_table_row(line: str) -> List[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{2,}", line))


def render_markdown_to_document(
    doc: Document,
    md_text: str,
    *,
    assets_dir: Optional[Path] = None,
) -> None:
    """Append markdown content to an existing document."""
    from src.modules.export.mermaid_renderer import add_mermaid_to_document
    from src.modules.generation.rewrite_validation import (
        dedupe_consecutive_section_headings,
        strip_section_id_tags,
    )

    md_text = dedupe_consecutive_section_headings(strip_section_id_tags(md_text))
    md_text = _normalize_notes_body_for_export(md_text)

    center_mode = False
    cover_band_added = False
    in_code_block = False
    mermaid_block = False
    mermaid_lines: List[str] = []
    table_rows: List[List[str]] = []
    list_tracker = NumberedListTracker()

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        while table_rows and all(not (c or "").strip() for c in table_rows[0]):
            table_rows.pop(0)
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for ri, row in enumerate(table_rows):
            for ci in range(cols):
                val = row[ci] if ci < len(row) else ""
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_inline_runs(p, val)
        is_metadata = (
            len(table_rows) >= 1
            and len(table_rows[0]) >= 2
            and table_rows[0][0].strip().lower() in {"", "field"}
            and table_rows[0][1].strip().lower() in {"", "value"}
        )
        if is_metadata:
            format_metadata_table(table)
        else:
            format_content_table(table)
        table_rows = []

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if in_code_block:
            if line.strip() == "```":
                if mermaid_block and mermaid_lines:
                    add_mermaid_to_document(doc, "\n".join(mermaid_lines), assets_dir=assets_dir)
                in_code_block = False
                mermaid_block = False
                mermaid_lines = []
            elif mermaid_block:
                mermaid_lines.append(line)
            elif _OPENXML_PAGE_BREAK.match(line):
                flush_table()
                _add_page_break(doc)
            continue

        if line.strip().startswith("```"):
            flush_table()
            fence = line.strip()[3:].strip().lower()
            if fence == "mermaid":
                in_code_block = True
                mermaid_block = True
                mermaid_lines = []
            elif fence.startswith("{=openxml}"):
                in_code_block = True
                mermaid_block = False
            continue

        if line.strip().startswith("|"):
            if _is_table_separator(line):
                continue
            table_rows.append(_parse_table_row(line))
            continue
        flush_table()

        if line.startswith("<div"):
            center_mode = "center" in line.lower()
            if center_mode and not cover_band_added:
                add_cover_header_band(doc)
                cover_band_added = True
            continue
        if line.startswith("</div"):
            center_mode = False
            continue

        if not line.strip():
            continue

        if line.startswith("### "):
            list_tracker.on_section_boundary()
            h = doc.add_heading(line[4:].strip(), level=3)
            style_section_heading(h, level=3)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("## "):
            list_tracker.on_section_boundary()
            h = doc.add_heading(line[3:].strip(), level=2)
            style_section_heading(h, level=2)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("# "):
            list_tracker.on_section_boundary()
            title_text = line[2:].strip()
            if title_text.lower() == "table of contents":
                title_p = doc.add_paragraph()
                tr = title_p.add_run(title_text)
                tr.bold = True
                style_toc_title(title_p)
                continue
            h = doc.add_heading(title_text, level=1)
            style_chapter_heading(h)
            if center_mode:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style_cover_title(h)
            continue

        if is_callout_label(line):
            list_tracker.on_non_ordered_line()
            label = callout_label_text(line)
            if is_named_callout_label(label):
                add_callout_paragraph(doc, "", label=label)
            else:
                add_topic_subheading(doc, label)
            continue

        if line.startswith("> "):
            list_tracker.on_non_ordered_line()
            add_callout_paragraph(doc, line[2:].strip())
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            list_tracker.on_non_ordered_line()
            indent, text = bullet.groups()
            level = min(len(indent.expandtabs()) // 2, 2)
            p = doc.add_paragraph(style=_bullet_style(level))
            _add_inline_runs(p, text.strip())
            style_bullet_paragraph(p, level=level)
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            indent, text = ordered.groups()
            level = min(len(indent.expandtabs()) // 2, 2)
            style = "List Number" if level == 0 else f"List Number {level + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Number")
            list_tracker.apply_restart_if_needed(p, level=level)
            _add_inline_runs(p, text.strip())
            style_numbered_paragraph(p, level=level)
            continue

        list_tracker.on_non_ordered_line()
        p = doc.add_paragraph()
        if center_mode:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line.startswith("#"):
                for run in p.runs:
                    run.bold = True
                style_cover_title(p)
            elif line.startswith("###"):
                style_cover_subtitle(p)
        _add_inline_runs(p, line.strip())
        if not center_mode:
            style_body_paragraph(p)

    flush_table()


def export_markdown_file_to_docx(
    md_text: str,
    output_path: str | Path,
    *,
    reference_docx: Optional[str] = None,
    theme: Optional[str] = None,
) -> str:
    """Convert a full notes .md file to .docx preserving its structure."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    assets_dir = out.parent / f"{out.stem}_diagrams"

    doc_title = ""
    for raw in md_text.splitlines():
        s = raw.strip()
        if s.startswith("# ") and not s.startswith("## "):
            doc_title = s[2:].strip()
            if doc_title.lower() != "table of contents":
                break

    if reference_docx and Path(reference_docx).exists():
        doc = Document(reference_docx)
        _clear_document_body(doc)
    else:
        doc = Document()

    docx_theme = resolve_docx_theme(theme)
    apply_study_notes_theme(doc, doc_title=doc_title or "Study Notes", theme=docx_theme)
    render_markdown_to_document(doc, md_text, assets_dir=assets_dir)
    finalize_word_document(doc)
    doc.save(str(out))
    refresh_word_fields(str(out))
    logger.info("Exported markdown to Word: %s", out)
    return str(out)
