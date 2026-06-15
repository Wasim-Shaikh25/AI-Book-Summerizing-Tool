"""Export structured study notes to Word with cover, TOC, and page breaks."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.modules.export.document_formatter import BookCoverMeta
from src.modules.export.docx_theme import (
    add_callout_paragraph,
    add_cover_header_band,
    apply_study_notes_theme,
    resolve_docx_theme,
    callout_label_text,
    finalize_word_document,
    format_metadata_table,
    insert_word_field,
    is_callout_label,
    refresh_word_fields,
    style_body_paragraph,
    style_chapter_heading,
    style_cover_subtitle,
    style_cover_title,
    style_section_heading,
    style_toc_entry,
    style_toc_title,
)
from src.modules.generation.rewrite_validation import (
    SECTION_ID_TAG,
    heading_similarity,
    normalize_heading,
    strip_redundant_section_heading,
    strip_section_id_tags,
)

logger = logging.getLogger(__name__)

_TOC_TAB_POSITION = Inches(6.2)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


class _BookmarkIds:
    def __init__(self) -> None:
        self._n = 0

    def next(self) -> int:
        self._n += 1
        return self._n


def _add_bookmark(paragraph, name: str, bookmark_ids: _BookmarkIds) -> None:
    """Wrap the first run of a heading paragraph with a named bookmark."""
    if not paragraph.runs:
        paragraph.add_run()
    run = paragraph.runs[0]
    bid = bookmark_ids.next()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    run._r.addprevious(start)
    run._r.addnext(end)


def _sanitize_bookmark_name(raw: str) -> str:
    name = re.sub(r"[^A-Za-z0-9_]", "_", raw.strip())[:40] or "item"
    if name[0].isdigit():
        name = f"b_{name}"
    return name


def _add_toc_line(doc: Document, *, level: int, title: str, bookmark: str) -> None:
    """One TOC row: title ........ page (PAGEREF to heading bookmark)."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.tab_stops.add_tab_stop(_TOC_TAB_POSITION, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if level >= 2:
        fmt.left_indent = Inches(0.4)
    label = p.add_run(title[:110])
    if level == 1:
        label.bold = True
    p.add_run("\t")
    insert_word_field(p, f"PAGEREF {bookmark} \\h")
    style_toc_entry(p, level=level)


def _add_toc_block(doc: Document, toc_rows: Sequence[tuple[int, str, str]]) -> None:
    """Add TOC page (title + PAGEREF lines + page break) at the current document end."""
    title = doc.add_paragraph()
    tr = title.add_run("Table of Contents")
    tr.bold = True
    style_toc_title(title)

    doc.add_paragraph()

    for level, heading, bookmark in toc_rows:
        _add_toc_line(doc, level=level, title=heading, bookmark=bookmark)

    _add_page_break(doc)


def _add_inline_runs(paragraph, text: str) -> None:
    """Add paragraph runs with basic **bold** / *italic* markup."""
    if not text:
        return
    parts: List[tuple[str, bool, bool]] = [(text, False, False)]
    for pattern, bold, italic in ((_BOLD_RE, True, False), (_ITALIC_RE, False, True)):
        new_parts: List[tuple[str, bool, bool]] = []
        for chunk, is_b, is_i in parts:
            if is_b or is_i:
                new_parts.append((chunk, is_b, is_i))
                continue
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    new_parts.append((chunk[pos : m.start()], False, False))
                new_parts.append((m.group(1), bold, italic))
                pos = m.end()
            if pos < len(chunk):
                new_parts.append((chunk[pos:], False, False))
        parts = new_parts or parts
    for chunk, is_b, is_i in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = is_b
        run.italic = is_i


def append_markdown_body(
    doc: Document,
    text: str,
    *,
    compact: bool = False,
    assets_dir: Optional[Path] = None,
    user_instruction: str = "",
) -> None:
    """Render note body markdown (bullets, subtopics, lists) into the document."""
    from src.modules.export.note_body_docx import append_note_body_markdown

    if not (text or "").strip():
        return
    append_note_body_markdown(
        doc,
        text,
        compact=compact,
        user_instruction=user_instruction,
    )


def _add_cover_page(doc: Document, cover: BookCoverMeta) -> None:
    title = (cover.title or "Study Notes").strip()
    add_cover_header_band(doc)

    t = doc.add_paragraph()
    run = t.add_run(title)
    run.bold = True
    style_cover_title(t)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(cover.subtitle or "Study Notes")
    style_cover_subtitle(sub)

    doc.add_paragraph()
    rows = [
        ("Book", title),
        ("Source PDF", cover.source_pdf),
        ("Generated", cover.generated_at),
        ("Chapters", str(cover.chapter_count) if cover.chapter_count else ""),
        ("Sections", str(cover.section_count) if cover.section_count else ""),
        ("Notes style", cover.user_instruction),
    ]
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for label, val in rows:
        if not (val or "").strip():
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)
    format_metadata_table(table)
    doc.add_paragraph()


def _add_toc_page(
    doc: Document,
    hierarchy: Dict[str, Any],
    *,
    toc_entries: Optional[Sequence[Any]] = None,
) -> None:
    """Deprecated — TOC is built via _add_toc_block in final document order."""
    return


def _collect_chapter_export_plan(
    hierarchy: Dict[str, Any],
    rewritten: Dict[str, str],
    *,
    bundle_size: int,
    bundle_export: bool,
    compact_toc: bool,
) -> tuple[List[tuple[int, str, str]], List[Dict[str, Any]]]:
    """Return (toc_rows, chapter_blocks) without mutating a Document."""
    from src.modules.generation.section_bundler import build_rewrite_bundles

    toc_rows: List[tuple[int, str, str]] = []
    chapter_blocks: List[Dict[str, Any]] = []
    use_bundles = bundle_export and bundle_size > 1
    toc_sections = not compact_toc and not use_bundles

    for ch in hierarchy.get("chapters") or []:
        ch_heading = str(ch.get("heading") or "").strip()
        if not ch_heading:
            continue
        sec_rows: List[Dict[str, Any]] = []
        for sec in ch.get("sections") or []:
            sid = str(sec.get("section_id") or "")
            body = (rewritten.get(sid) or "").strip()
            if not body:
                continue
            sec_rows.append(
                {
                    "section_id": sid,
                    "heading": str(sec.get("heading") or sid).strip(),
                    "chapter_heading": ch_heading,
                    "text": body,
                }
            )
        if not sec_rows:
            continue

        ch_id = str(ch.get("chapter_id") or ch_heading)
        ch_bm = _sanitize_bookmark_name(f"ch_{ch_id}")
        toc_rows.append((1, ch_heading, ch_bm))

        section_entries: List[Dict[str, Any]] = []
        if use_bundles:
            for bundle in build_rewrite_bundles(sec_rows, bundle_size=bundle_size):
                bundle_bm = _sanitize_bookmark_name(f"b_{bundle.bundle_id}_{ch_id}")
                if toc_sections:
                    toc_rows.append((2, bundle.label[:110], bundle_bm))
                bundle_sections: List[Dict[str, Any]] = []
                for sid, sec_heading in zip(bundle.section_ids, bundle.headings):
                    body = (rewritten.get(sid) or "").strip()
                    if not body:
                        continue
                    sec_bm = _sanitize_bookmark_name(f"sec_{sid}")
                    if toc_sections:
                        toc_rows.append((3, sec_heading[:110], sec_bm))
                    bundle_sections.append(
                        {
                            "bookmark": sec_bm,
                            "heading": sec_heading,
                            "level": 3,
                            "text": strip_redundant_section_heading(body, sec_heading),
                            "compact": True,
                        }
                    )
                section_entries.append(
                    {
                        "kind": "bundle",
                        "bookmark": bundle_bm,
                        "heading": bundle.label,
                        "level": 2,
                        "sections": bundle_sections,
                    }
                )
        else:
            for row in sec_rows:
                sid = str(row.get("section_id") or "")
                sec_heading = str(row.get("heading") or sid)
                sec_bm = _sanitize_bookmark_name(f"sec_{sid}")
                if toc_sections:
                    toc_rows.append((2, sec_heading[:110], sec_bm))
                section_entries.append(
                    {
                        "kind": "section",
                        "bookmark": sec_bm,
                        "heading": sec_heading,
                        "level": 2,
                        "text": strip_redundant_section_heading(
                            str(row.get("text") or ""),
                            sec_heading,
                        ),
                        "compact": compact_toc,
                    }
                )

        chapter_blocks.append(
            {
                "bookmark": ch_bm,
                "heading": ch_heading,
                "sections": section_entries,
            }
        )

    return toc_rows, chapter_blocks


def _append_chapter_blocks(
    doc: Document,
    chapter_blocks: Sequence[Dict[str, Any]],
    *,
    bookmark_ids: _BookmarkIds,
    chapter_page_breaks: bool,
    assets_dir: Optional[Path] = None,
    user_instruction: str = "",
) -> None:
    for chapter_index, block in enumerate(chapter_blocks):
        if chapter_index > 0 and chapter_page_breaks:
            _add_page_break(doc)

        h1 = doc.add_heading(str(block.get("heading") or ""), level=1)
        style_chapter_heading(h1)
        _add_bookmark(h1, str(block["bookmark"]), bookmark_ids)

        for entry in block.get("sections") or []:
            if entry.get("kind") == "bundle":
                h2 = doc.add_heading(str(entry.get("heading") or "")[:200], level=2)
                style_section_heading(h2, level=2)
                _add_bookmark(h2, str(entry["bookmark"]), bookmark_ids)
                for sec in entry.get("sections") or []:
                    h3 = doc.add_heading(str(sec.get("heading") or "")[:200], level=3)
                    style_section_heading(h3, level=3)
                    _add_bookmark(h3, str(sec["bookmark"]), bookmark_ids)
                    append_markdown_body(
                        doc,
                        str(sec.get("text") or ""),
                        compact=True,
                        assets_dir=assets_dir,
                        user_instruction=user_instruction,
                    )
                continue

            level = int(entry.get("level") or 2)
            heading = doc.add_heading(str(entry.get("heading") or "")[:200], level=level)
            style_section_heading(heading, level=level)
            _add_bookmark(heading, str(entry["bookmark"]), bookmark_ids)
            append_markdown_body(
                doc,
                str(entry.get("text") or ""),
                compact=bool(entry.get("compact")),
                assets_dir=assets_dir,
                user_instruction=user_instruction,
            )


class DocxNotesExporter:
    """Build a formatted Word document from 15e hierarchy + rewritten section bodies."""

    def export(
        self,
        output_path: str | Path,
        *,
        cover: BookCoverMeta,
        hierarchy: Dict[str, Any],
        rewritten: Dict[str, str],
        reference_docx: Optional[str] = None,
        bundle_size: int = 1,
        bundle_export: bool = False,
        compact_toc: bool = False,
        chapter_page_breaks: Optional[bool] = None,
    ) -> str:
        from src.modules.generation.section_bundler import resolve_chapter_page_breaks

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        use_bundles = bundle_export and bundle_size > 1
        if chapter_page_breaks is None:
            chapter_page_breaks = resolve_chapter_page_breaks(
                compact_toc=compact_toc,
                use_bundles=use_bundles,
            )

        toc_rows, chapter_blocks = _collect_chapter_export_plan(
            hierarchy,
            rewritten,
            bundle_size=bundle_size,
            bundle_export=bundle_export,
            compact_toc=compact_toc,
        )
        if not chapter_blocks:
            raise ValueError("No chapter content available for Word export")

        if reference_docx and Path(reference_docx).exists():
            doc = Document(reference_docx)
        else:
            doc = Document()
        apply_study_notes_theme(
            doc,
            doc_title=(cover.title or "Study Notes").strip(),
            theme=resolve_docx_theme(),
        )
        bookmark_ids = _BookmarkIds()
        assets_dir = out.parent / f"{out.stem}_diagrams"

        # Final order: cover -> TOC -> chapters. Avoid prepending TOC after body build,
        # which leaves stale PAGEREF page numbers until Word refreshes all fields.
        _add_cover_page(doc, cover)
        _add_page_break(doc)
        if toc_rows:
            _add_toc_block(doc, toc_rows)
        _append_chapter_blocks(
            doc,
            chapter_blocks,
            bookmark_ids=bookmark_ids,
            chapter_page_breaks=chapter_page_breaks,
            assets_dir=assets_dir,
            user_instruction=(cover.user_instruction or "").strip(),
        )

        finalize_word_document(doc)
        doc.save(str(out))
        refresh_word_fields(str(out))
        logger.info("Saved formatted Word document: %s", out)
        return str(out)


def parse_markdown_sections(md_text: str) -> Tuple[Dict[str, str], Dict[str, str]]:
    """Return (by_heading, by_section_id) from markdown body."""
    m = re.search(r"# Table of Contents[\s\S]*?\n# ([^\n]+)", md_text)
    body = md_text[m.start(1) :] if m else md_text

    by_heading: Dict[str, str] = {}
    by_sid: Dict[str, str] = {}
    current_heading = ""
    current_sid = ""
    buf: List[str] = []
    skip_block = False

    def flush() -> None:
        nonlocal buf, current_heading, current_sid
        if not buf:
            return
        text = "\n".join(buf).strip()
        if not text:
            buf = []
            return
        if current_sid:
            by_sid[current_sid] = text
        if current_heading:
            by_heading[current_heading] = text
        buf = []

    for line in body.splitlines():
        if line.strip().startswith("```{=openxml}"):
            skip_block = True
            continue
        if skip_block:
            if line.strip() == "```":
                skip_block = False
            continue
        if line.startswith("# ") and not line.startswith("## "):
            flush()
            current_heading = ""
            current_sid = ""
            continue
        if line.startswith("## ") and not line.startswith("### "):
            flush()
            raw = line[3:].strip()
            sid_m = SECTION_ID_TAG.search(raw)
            current_sid = sid_m.group(1) if sid_m else ""
            current_heading = SECTION_ID_TAG.sub("", raw).strip()
            continue
        if line.startswith("### ") and not line.startswith("#### "):
            flush()
            raw = line[4:].strip()
            sid_m = SECTION_ID_TAG.search(raw)
            current_sid = sid_m.group(1) if sid_m else ""
            current_heading = SECTION_ID_TAG.sub("", raw).strip()
            continue
        if current_heading or current_sid:
            buf.append(line)
    flush()
    return by_heading, by_sid


def parse_section_bodies_from_markdown(md_text: str) -> Dict[str, str]:
    """Map section heading -> body text (backward compatible)."""
    by_heading, _ = parse_markdown_sections(md_text)
    return by_heading


def _norm_heading(text: str) -> str:
    return normalize_heading(text)


def _match_body_for_section(
    *,
    sid: str,
    heading: str,
    by_sid: Dict[str, str],
    by_heading: Dict[str, str],
    used_headings: set[str],
    fuzzy_threshold: float,
) -> str:
    if sid and sid in by_sid:
        return by_sid[sid]
    if heading in by_heading and heading not in used_headings:
        used_headings.add(heading)
        return by_heading[heading]
    from src.modules.structure.final_structuring.heading_cleanup import canonical_heading_for_match

    nh = canonical_heading_for_match(heading)
    for key, val in by_heading.items():
        if key in used_headings:
            continue
        if canonical_heading_for_match(key) == nh:
            used_headings.add(key)
            return val
    best_key = ""
    best_score = 0.0
    for key, val in by_heading.items():
        if key in used_headings:
            continue
        score = heading_similarity(canonical_heading_for_match(key), nh)
        if score > best_score:
            best_score = score
            best_key = key
    if best_score >= fuzzy_threshold and best_key:
        used_headings.add(best_key)
        return by_heading[best_key]
    return ""


def rewritten_map_from_section_bodies(
    hierarchy: Dict[str, Any],
    section_bodies: Dict[str, str],
    *,
    by_section_id: Optional[Dict[str, str]] = None,
    md_text: Optional[str] = None,
    fuzzy_threshold: float = 0.82,
) -> Dict[str, str]:
    """Match section bodies to section_id by sid tag, exact heading, then fuzzy match."""
    by_heading = dict(section_bodies)
    by_sid = dict(by_section_id or {})
    used_headings: set[str] = set()
    result: Dict[str, str] = {}

    for ch in hierarchy.get("chapters") or []:
        for sec in ch.get("sections") or []:
            sid = str(sec.get("section_id") or "")
            heading = str(sec.get("heading") or "").strip()
            body = _match_body_for_section(
                sid=sid,
                heading=heading,
                by_sid=by_sid,
                by_heading=by_heading,
                used_headings=used_headings,
                fuzzy_threshold=fuzzy_threshold,
            )
            if body:
                result[sid] = body

    # Positional fallback per chapter when MD section count matches hierarchy
    md_by_chapter = _sections_by_chapter_from_markdown(md_text) if md_text else {}
    if not md_by_chapter and by_heading:
        md_by_chapter = {"": [(k, v) for k, v in by_heading.items()]}
    for ch in hierarchy.get("chapters") or []:
        ch_name = _norm_heading(str(ch.get("heading") or ""))
        md_secs = md_by_chapter.get(ch_name) or []
        h_secs = [s for s in ch.get("sections") or [] if str(s.get("section_id") or "") not in result]
        if not h_secs or not md_secs:
            continue
        if len(md_secs) == len(h_secs):
            for sec, (_, body) in zip(h_secs, md_secs):
                sid = str(sec.get("section_id") or "")
                if sid and body:
                    result[sid] = body
    return result


def _sections_by_chapter_from_markdown(md_text: str) -> Dict[str, List[Tuple[str, str]]]:
    """Group ## sections under their preceding # chapter heading."""
    m = re.search(r"# Table of Contents[\s\S]*?\n# ([^\n]+)", md_text)
    body = md_text[m.start(1) :] if m else md_text
    out: Dict[str, List[Tuple[str, str]]] = {}
    chapter_key = ""
    current_heading = ""
    buf: List[str] = []
    skip_block = False

    def flush() -> None:
        nonlocal buf, current_heading
        if current_heading and buf:
            text = "\n".join(buf).strip()
            if text:
                out.setdefault(chapter_key, []).append((current_heading, text))
        buf = []

    for line in body.splitlines():
        if line.strip().startswith("```{=openxml}"):
            skip_block = True
            continue
        if skip_block:
            if line.strip() == "```":
                skip_block = False
            continue
        if line.startswith("# ") and not line.startswith("## "):
            flush()
            chapter_key = _norm_heading(line[2:].strip())
            current_heading = ""
            continue
        if line.startswith("## ") and not line.startswith("### "):
            flush()
            raw = line[3:].strip()
            current_heading = SECTION_ID_TAG.sub("", raw).strip()
            continue
        if current_heading:
            buf.append(line)
    flush()
    return out


def resolve_rewritten_map(
    hierarchy: Dict[str, Any],
    *,
    rewritten_by_id: Optional[Dict[str, str]] = None,
    md_text: Optional[str] = None,
) -> Dict[str, str]:
    """Prefer section_id sidecar; fall back to markdown parsing."""
    if rewritten_by_id:
        return {str(k): str(v) for k, v in rewritten_by_id.items() if v and str(v).strip()}
    if md_text:
        by_heading, by_sid = parse_markdown_sections(md_text)
        if by_sid:
            result: Dict[str, str] = {}
            for ch in hierarchy.get("chapters") or []:
                for sec in ch.get("sections") or []:
                    sid = str(sec.get("section_id") or "")
                    if sid and sid in by_sid:
                        result[sid] = by_sid[sid]
            if len(result) >= len(by_sid) * 0.5:
                return result
        return rewritten_map_from_section_bodies(
            hierarchy, by_heading, by_section_id=by_sid, md_text=md_text
        )
    return {}
