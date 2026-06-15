"""Render mermaid diagram blocks to PNG for Word export."""

from __future__ import annotations

import hashlib
import logging
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, List, Optional, Sequence, Tuple
from urllib import request

logger = logging.getLogger(__name__)

_MERMAID_FENCE_RE = re.compile(r"```mermaid\s*\n([\s\S]*?)\n```", re.I)
_STANDALONE_DIAGRAM_HEADING_RE = re.compile(r"^#{1,3}\s*diagram\s*$", re.I)
_NODE_RE = re.compile(
    r"([A-Za-z][\w]*)\s*(?:\[\s*([^\]]+)\]|\(\s*([^)]+)\)|\{\s*([^}]+)\})",
)
_EDGE_RE = re.compile(
    r"([A-Za-z][\w]*)(?:\[[^\]]*\]|\([^)]*\)|\{[^}]*\})?"
    r"\s*(?:-->|---|-.->|==>)\s*(?:\|[^|]*\|\s*)?"
    r"([A-Za-z][\w]*)(?:\[[^\]]*\]|\([^)]*\)|\{[^}]*\})?",
)

_DEFAULT_MAX_WIDTH_IN = 5.25
_DEFAULT_MAX_HEIGHT_IN = 3.75
_DEFAULT_MMDC_WIDTH = 680
_DEFAULT_MMDC_SCALE = 0.92
_PAGE_CONTENT_WIDTH_IN = 6.0
_READABLE_DPI = 105.0


@dataclass(frozen=True)
class MarkdownSegment:
    kind: str  # "markdown" | "mermaid"
    content: str


def _env_float(key: str, default: float) -> float:
    raw = (os.environ.get(key) or "").strip()
    if not raw:
        return default
    try:
        return float(raw)
    except ValueError:
        return default


def diagram_max_width_inches() -> float:
    return _env_float("DIAGRAM_MAX_WIDTH_IN", _DEFAULT_MAX_WIDTH_IN)


def diagram_max_height_inches() -> float:
    return _env_float("DIAGRAM_MAX_HEIGHT_IN", _DEFAULT_MAX_HEIGHT_IN)


def _strip_diagram_heading_markdown(text: str) -> str:
    """Remove standalone '### Diagram' headings — image is self-explanatory in Word."""
    lines: List[str] = []
    for line in (text or "").splitlines():
        if _STANDALONE_DIAGRAM_HEADING_RE.match(line.strip()):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def iter_markdown_segments(text: str) -> Iterator[MarkdownSegment]:
    """Split markdown into alternating prose and mermaid diagram blocks."""
    raw = text or ""
    if not raw.strip():
        return
    pos = 0
    for match in _MERMAID_FENCE_RE.finditer(raw):
        if match.start() > pos:
            chunk = _strip_diagram_heading_markdown(raw[pos : match.start()])
            if chunk:
                yield MarkdownSegment("markdown", chunk)
        body = (match.group(1) or "").strip()
        if body:
            yield MarkdownSegment("mermaid", body)
        pos = match.end()
    tail = _strip_diagram_heading_markdown(raw[pos:])
    if tail:
        yield MarkdownSegment("markdown", tail)


def _cache_path(assets_dir: Path, mermaid_src: str) -> Path:
    digest = hashlib.sha1(mermaid_src.encode("utf-8")).hexdigest()[:16]
    assets_dir.mkdir(parents=True, exist_ok=True)
    return assets_dir / f"diagram_v3_{digest}.png"


def _run_mmdc(mermaid_src: str, out_path: Path) -> bool:
    runners = []
    if shutil.which("mmdc"):
        runners.append(["mmdc"])
    if shutil.which("npx"):
        runners.append(["npx", "-y", "@mermaid-js/mermaid-cli"])
    width = int(_env_float("MMDC_WIDTH", _DEFAULT_MMDC_WIDTH))
    scale = _env_float("MMDC_SCALE", _DEFAULT_MMDC_SCALE)
    for base in runners:
        tmp_path = ""
        try:
            with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False, encoding="utf-8") as tmp:
                tmp.write(mermaid_src)
                tmp_path = tmp.name
            cmd = [
                *base,
                "-i",
                tmp_path,
                "-o",
                str(out_path),
                "-b",
                "white",
                "-w",
                str(width),
                "-s",
                str(scale),
            ]
            subprocess.run(cmd, capture_output=True, text=True, check=True, timeout=90)
            Path(tmp_path).unlink(missing_ok=True)
            if out_path.exists() and out_path.stat().st_size > 0:
                return True
        except Exception as exc:
            logger.debug("mermaid CLI %s failed: %s", base[0], exc)
            if tmp_path:
                try:
                    Path(tmp_path).unlink(missing_ok=True)
                except Exception:
                    pass
    return False


def _run_kroki(mermaid_src: str, out_path: Path) -> bool:
    try:
        req = request.Request(
            "https://kroki.io/mermaid/png",
            data=mermaid_src.encode("utf-8"),
            headers={"Content-Type": "text/plain"},
            method="POST",
        )
        with request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        if data:
            out_path.write_bytes(data)
            _shrink_png_if_needed(out_path)
            return True
    except Exception as exc:
        logger.debug("Kroki mermaid render failed: %s", exc)
    return False


def _wrap_text(text: str, *, max_chars: int = 14) -> List[str]:
    words = (text or "").split()
    if not words:
        return [""]
    lines: List[str] = []
    current: List[str] = []
    length = 0
    for word in words:
        add = len(word) + (1 if current else 0)
        if current and length + add > max_chars:
            lines.append(" ".join(current))
            current = [word]
            length = len(word)
        else:
            current.append(word)
            length += add
    if current:
        lines.append(" ".join(current))
    return lines[:3]


def _parse_flowchart(mermaid_src: str) -> Tuple[List[Tuple[str, str]], List[Tuple[str, str]]]:
    nodes: dict[str, str] = {}
    edges: List[Tuple[str, str]] = []
    for line in mermaid_src.splitlines():
        s = line.strip()
        if not s or s.lower().startswith(("flowchart", "graph", "%%")):
            continue
        for m in _EDGE_RE.finditer(s):
            a, b = m.group(1), m.group(2)
            nodes.setdefault(a, a)
            nodes.setdefault(b, b)
            edges.append((a, b))
        for m in _NODE_RE.finditer(s):
            nid = m.group(1)
            label = next((g for g in m.groups()[1:] if g), nid) or nid
            nodes[nid] = label.strip()
    if not nodes and not edges:
        return [], []
    ordered: List[Tuple[str, str]] = []
    seen: set[str] = set()
    for a, b in edges:
        for nid in (a, b):
            if nid not in seen:
                ordered.append((nid, nodes.get(nid, nid)))
                seen.add(nid)
    for nid, label in nodes.items():
        if nid not in seen:
            ordered.append((nid, label))
            seen.add(nid)
    return ordered, edges


def _render_with_pillow(mermaid_src: str, out_path: Path) -> bool:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return False

    nodes, edges = _parse_flowchart(mermaid_src)
    if not nodes:
        return False

    node_w, node_h, gap_y, gap_x = 150, 48, 20, 24
    cols = 1 if len(nodes) <= 5 else 2
    rows = (len(nodes) + cols - 1) // cols
    width = cols * node_w + (cols + 1) * gap_x
    height = rows * node_h + (rows + 1) * gap_y + 12

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 11)
    except Exception:
        font = ImageFont.load_default()

    positions: dict[str, Tuple[int, int, int, int]] = {}
    for idx, (nid, label) in enumerate(nodes):
        col = idx % cols
        row = idx // cols
        x0 = gap_x + col * (node_w + gap_x)
        y0 = gap_y + row * (node_h + gap_y)
        x1, y1 = x0 + node_w, y0 + node_h
        positions[nid] = (x0, y0, x1, y1)
        draw.rounded_rectangle((x0, y0, x1, y1), radius=8, outline="#1f4e79", width=1, fill="#eef4fb")
        wrapped = _wrap_text(label)
        ty = y0 + 8
        for wline in wrapped:
            draw.text((x0 + 8, ty), wline, fill="#102030", font=font)
            ty += 14

    for a, b in edges:
        if a not in positions or b not in positions:
            continue
        ax0, ay0, ax1, ay1 = positions[a]
        bx0, by0, bx1, by1 = positions[b]
        start = ((ax0 + ax1) // 2, ay1)
        end = ((bx0 + bx1) // 2, by0)
        draw.line([start, end], fill="#444444", width=1)
        draw.polygon(
            [
                (end[0], end[1]),
                (end[0] - 4, end[1] - 7),
                (end[0] + 4, end[1] - 7),
            ],
            fill="#444444",
        )

    img.save(out_path, format="PNG")
    return out_path.exists()


def _shrink_png_if_needed(png_path: Path) -> None:
    """Downscale very large PNGs so Word layout stays compact."""
    try:
        from PIL import Image
    except ImportError:
        return
    max_px = int(_env_float("DIAGRAM_MAX_PX", 1100))
    try:
        with Image.open(png_path) as im:
            w, h = im.size
            if max(w, h) <= max_px:
                return
            scale = max_px / float(max(w, h))
            new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
            resized = im.resize(new_size, Image.Resampling.LANCZOS)
            resized.save(png_path, format="PNG")
    except Exception as exc:
        logger.debug("PNG shrink skipped for %s: %s", png_path, exc)


def render_mermaid_to_png(
    mermaid_src: str,
    out_path: Path,
    *,
    allow_network: bool = True,
) -> bool:
    """Render mermaid source to PNG using CLI, Kroki, or Pillow fallback."""
    src = (mermaid_src or "").strip()
    if not src:
        return False
    out_path.parent.mkdir(parents=True, exist_ok=True)

    ok = False
    if _run_mmdc(src, out_path):
        ok = True
    elif allow_network and _run_kroki(src, out_path):
        ok = True
    elif _render_with_pillow(src, out_path):
        ok = True
    if ok:
        _shrink_png_if_needed(out_path)
    return ok


def render_mermaid_cached(
    mermaid_src: str,
    assets_dir: Path,
    *,
    allow_network: bool = True,
) -> Optional[Path]:
    """Return cached PNG path for a mermaid block."""
    out_path = _cache_path(assets_dir, mermaid_src)
    if out_path.exists() and out_path.stat().st_size > 0:
        return out_path
    if render_mermaid_to_png(mermaid_src, out_path, allow_network=allow_network):
        return out_path
    return None


def _fit_picture_dimensions(
    png_path: Path,
    *,
    max_width_in: float,
    max_height_in: float,
) -> Tuple[float, float]:
    """Return width/height in inches, preserving aspect ratio within bounds."""
    try:
        from PIL import Image
    except ImportError:
        return max_width_in, max_height_in * 0.6

    with Image.open(png_path) as im:
        px_w, px_h = im.size
    if px_w <= 0 or px_h <= 0:
        return max_width_in, max_height_in * 0.5
    aspect = px_h / px_w
    width = max_width_in
    height = width * aspect
    if height > max_height_in:
        height = max_height_in
        width = height / aspect
    return width, height


def _diagram_complexity_caps(mermaid_src: str) -> Tuple[float, float, bool]:
    """Return (max_width_in, max_height_in, is_horizontal) from graph shape."""
    nodes, _ = _parse_flowchart(mermaid_src)
    node_count = max(len(nodes), 1)
    is_horizontal = bool(re.search(r"(?:flowchart|graph)\s+LR\b", mermaid_src, re.I))

    if node_count <= 3:
        max_w, max_h = 4.75, 3.25
    elif node_count <= 5:
        max_w, max_h = 5.25, 3.75
    elif node_count <= 8:
        max_w, max_h = 5.75, 4.25
    else:
        max_w, max_h = min(_PAGE_CONTENT_WIDTH_IN, 6.25), 4.75

    if is_horizontal:
        max_w = min(_PAGE_CONTENT_WIDTH_IN, max_w + 0.65)
        max_h = min(4.0, max_h)
    return max_w, max_h, is_horizontal


def compute_diagram_display_size(
    png_path: Path,
    mermaid_src: str,
    *,
    page_content_width_in: float = _PAGE_CONTENT_WIDTH_IN,
) -> Tuple[float, float]:
    """
    Size diagrams for readability.

    Uses node count for complexity caps and pixel dimensions for a readable DPI floor.
    """
    max_w, max_h, _ = _diagram_complexity_caps(mermaid_src)
    max_w = min(max_w, page_content_width_in)

    try:
        from PIL import Image
    except ImportError:
        return _fit_picture_dimensions(png_path, max_width_in=max_w, max_height_in=max_h)

    with Image.open(png_path) as im:
        px_w, px_h = im.size
    if px_w <= 0 or px_h <= 0:
        return _fit_picture_dimensions(png_path, max_width_in=max_w, max_height_in=max_h)

    width, height = _fit_picture_dimensions(png_path, max_width_in=max_w, max_height_in=max_h)

    # Ensure labels stay legible: target ~105 DPI unless that exceeds page bounds.
    min_w = min(px_w / _READABLE_DPI, max_w)
    min_h = min(px_h / _READABLE_DPI, max_h)
    if width < min_w * 0.9:
        width = min_w
        height = width * (px_h / px_w)
        if height > max_h:
            height = max_h
            width = height * (px_w / px_h)
    if height < min_h * 0.9 and height < max_h:
        height = min(min_h, max_h)
        width = height * (px_w / px_h)
        if width > max_w:
            width = max_w
            height = width * (px_h / px_w)

    return width, height


def _section_body_height_inches(doc) -> float:
    """Approximate printable body height for the active section (inches)."""
    try:
        section = doc.sections[-1]
        page_h = section.page_height.inches
        top = section.top_margin.inches
        bottom = section.bottom_margin.inches
        return max(5.0, page_h - top - bottom - 0.45)
    except Exception:
        return 8.0


def shrink_diagram_to_page_flow(
    width_in: float,
    height_in: float,
    *,
    page_body_height_in: float,
    reserved_above_in: float = 0.55,
    min_width_in: float = 2.75,
) -> Tuple[float, float]:
    """
    Scale a diagram down so it flows with nearby text instead of orphaning a page.

    Caps height to the likely remaining space on a page (not the preferred large size).
    """
    if width_in <= 0 or height_in <= 0:
        return width_in, height_in

    flow_slot = max(2.0, page_body_height_in - reserved_above_in)
    page_cap = page_body_height_in * 0.68
    target_h = min(flow_slot, page_cap)
    if height_in <= target_h:
        return width_in, height_in

    scale = target_h / height_in
    new_w = max(min_width_in, width_in * scale)
    new_h = height_in * (new_w / width_in) if width_in else target_h
    if new_h > target_h:
        scale = target_h / new_h
        new_w = max(min_width_in, new_w * scale)
        new_h = target_h
    return new_w, new_h


def add_mermaid_to_document(doc, mermaid_src: str, *, assets_dir: Optional[Path] = None) -> bool:
    """Insert a compact rendered diagram image into a python-docx Document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    base_dir = assets_dir or Path(tempfile.gettempdir()) / "ai_notes_diagrams"
    png_path = render_mermaid_cached(mermaid_src, base_dir)
    if not png_path:
        p = doc.add_paragraph()
        run = p.add_run("Diagram (could not render)")
        run.italic = True
        return False

    width_in, height_in = compute_diagram_display_size(png_path, mermaid_src)
    body_h = _section_body_height_inches(doc)
    width_in, height_in = shrink_diagram_to_page_flow(
        width_in,
        height_in,
        page_body_height_in=body_h,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = False
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.widow_control = False
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(width_in), height=Inches(height_in))
    return True
