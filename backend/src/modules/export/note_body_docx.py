"""Line-by-line markdown note body → Word paragraphs (bullets, subtopics, lists)."""

from __future__ import annotations

import re
from typing import Optional

from docx import Document
from docx.shared import Pt

from src.modules.export.docx_theme import (
    NumberedListTracker,
    add_callout_paragraph,
    add_topic_subheading,
    callout_label_text,
    is_callout_label,
    is_named_callout_label,
    style_body_paragraph,
    style_bullet_paragraph,
    style_numbered_paragraph,
    style_section_heading,
)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
_STANDALONE_BOLD_RE = re.compile(r"^\*\*(.+?)\*\*:?\s*$")


def _bullet_style(level: int) -> str:
    if level <= 0:
        return "List Bullet"
    if level == 1:
        return "List Bullet 2"
    return "List Bullet 3"


def _add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return
    parts: list[tuple[str, bool, bool]] = [(text, False, False)]
    for pattern, bold, italic in ((_BOLD_RE, True, False), (_ITALIC_RE, False, True)):
        new_parts: list[tuple[str, bool, bool]] = []
        for chunk, is_b, is_i in parts:
            if is_b or is_i:
                new_parts.append((chunk, is_b, is_i))
                continue
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    new_parts.append((chunk[pos : m.start()], False, False))
                new_parts.append((m.group(1), bold, italic))
                pos = m.end()
            if pos < len(chunk):
                new_parts.append((chunk[pos:], False, False))
        parts = new_parts or parts
    for chunk, is_b, is_i in parts:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = is_b
        run.italic = is_i


def append_note_body_markdown(
    doc: Document,
    text: str,
    *,
    compact: bool = False,
    user_instruction: str = "",
) -> None:
    """Render note body markdown with proper bullets, numbered lists, and subtopic labels."""
    from src.modules.export.mermaid_renderer import add_mermaid_to_document, iter_markdown_segments
    from src.modules.generation.markdown_format_normalizer import strict_normalize_markdown
    from src.modules.generation.rewrite_validation import strip_section_id_tags

    if not (text or "").strip():
        return

    normalized = strict_normalize_markdown(text, user_instruction=user_instruction)
    space_after = Pt(2) if compact else None

    for segment in iter_markdown_segments(normalized):
        if segment.kind == "mermaid":
            add_mermaid_to_document(doc, segment.content)
            continue
        _append_plain_lines(doc, segment.content, compact=compact, space_after=space_after)


def _append_plain_lines(
    doc: Document,
    text: str,
    *,
    compact: bool,
    space_after: Optional[Pt],
) -> None:
    from src.modules.generation.rewrite_validation import strip_section_id_tags

    list_tracker = NumberedListTracker()

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            list_tracker.on_section_boundary()
            h = doc.add_heading(strip_section_id_tags(stripped[4:]).strip(), level=3)
            style_section_heading(h, level=3)
            if compact and space_after:
                h.paragraph_format.space_after = space_after
            continue

        if stripped.startswith("## "):
            list_tracker.on_section_boundary()
            h = doc.add_heading(strip_section_id_tags(stripped[3:]).strip(), level=2)
            style_section_heading(h, level=2)
            if compact and space_after:
                h.paragraph_format.space_after = space_after
            continue

        # Disabled internal subtitles for uniform structure - treat as regular text
        # if _STANDALONE_BOLD_RE.match(stripped) or is_callout_label(stripped):
        #     list_tracker.on_non_ordered_line()
        #     label = callout_label_text(stripped)
        #     try:
        #         from src.shared.notes_export_style import is_book_export_style

        #         book = is_book_export_style()
        #     except Exception:
        #         book = False
        #     if book:
        #         p = doc.add_paragraph()
        #         run = p.add_run(label.rstrip(":"))
        #         run.bold = True
        #         style_body_paragraph(p)
        #         if compact and space_after:
        #             p.paragraph_format.space_after = space_after
        #     elif is_named_callout_label(label):
        #         add_callout_paragraph(doc, "", label=label)
        #     else:
        #         add_topic_subheading(doc, label)
        #     continue

        # Disabled callout paragraphs for uniform structure - treat as regular text
        # if stripped.startswith("> "):
        #     list_tracker.on_non_ordered_line()
        #     add_callout_paragraph(doc, stripped[2:].strip())
        #     continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            list_tracker.on_non_ordered_line()
            indent, body = bullet.groups()
            level = min(len(indent.expandtabs()) // 2, 2)
            try:
                p = doc.add_paragraph(style=_bullet_style(level))
            except KeyError:
                p = doc.add_paragraph()
            _add_inline_runs(p, body.strip())
            style_bullet_paragraph(p, level=level)
            if compact and space_after:
                p.paragraph_format.space_after = space_after
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            indent, body = ordered.groups()
            level = min(len(indent.expandtabs()) // 2, 2)
            list_tracker.on_non_ordered_line()
            num_m = re.match(r"^(\s*)(\d+)\.", line)
            num = num_m.group(2) if num_m else "1"
            # Render as plain text "N. body" — not Word List Number. Inline
            # w:lvlOverride on numPr breaks Microsoft Word (invalid OOXML);
            # markdown already restarts numbering per section.
            p = doc.add_paragraph()
            prefix = p.add_run(f"{num}. ")
            prefix.bold = True
            _add_inline_runs(p, body.strip())
            style_bullet_paragraph(p, level=level)
            if compact and space_after:
                p.paragraph_format.space_after = space_after
            continue

        list_tracker.on_non_ordered_line()
        p = doc.add_paragraph()
        _add_inline_runs(p, strip_section_id_tags(stripped))
        style_body_paragraph(p)
        if compact and space_after:
            p.paragraph_format.space_after = space_after
