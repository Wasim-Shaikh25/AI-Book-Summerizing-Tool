"""Professional Word styling for study notes — color or black & white themes."""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from src.modules.export.docx_theme_palettes import (
    DocxThemePalette,
    get_palette,
    normalize_theme_name,
)
from src.shared.document_format_style import resolve_typography

_CALLOUT_LABEL_RE = re.compile(r"^\*\*(.+?)\*\*:?\s*$")

logger = logging.getLogger(__name__)

_active_palette: DocxThemePalette = get_palette("color")


def resolve_docx_theme(name: str | None = None) -> str:
    """Resolve theme from arg, env DOCX_THEME, or config (default: color)."""
    if name:
        return normalize_theme_name(name)
    raw = os.environ.get("DOCX_THEME", "").strip()
    if raw:
        return normalize_theme_name(raw)
    try:
        from src import config as cfg

        return normalize_theme_name(str(getattr(cfg, "DOCX_THEME", "color")))
    except Exception:
        return "color"


def _use_palette(theme: str | None = None) -> DocxThemePalette:
    global _active_palette
    _active_palette = get_palette(resolve_docx_theme(theme))
    return _active_palette


def _p() -> DocxThemePalette:
    return _active_palette


def _typo():
    return resolve_typography()


def _body_font() -> str:
    return _typo().body_font


def _heading_font() -> str:
    return _typo().heading_font


def _first_line_indent_enabled() -> bool:
    raw = os.environ.get("DOCX_FIRST_LINE_INDENT", "1").strip().lower()
    return raw not in {"0", "false", "no", "off"}


def _set_run_font(run, *, name: str | None = None, size: Optional[Pt] = None, color: Optional[RGBColor] = None) -> None:
    run.font.name = name or _body_font()
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def _set_paragraph_border(
    paragraph,
    *,
    color: str = "0078D4",
    size: int = 12,
    sides: tuple[str, ...] = ("bottom",),
) -> None:
    """Border on paragraph edges (size in eighths of a point)."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in sides:
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), color)
        p_bdr.append(edge)
    p_pr.append(p_bdr)


def _remove_table_borders(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def apply_study_notes_theme(
    doc: Document,
    *,
    doc_title: str = "",
    theme: str | None = None,
) -> None:
    """Apply fonts, spacing, margins, and heading colors (color or bw theme)."""
    palette = _use_palette(theme)
    typo = _typo()
    for section in doc.sections:
        section.top_margin = Inches(typo.margin_top_inches)
        section.bottom_margin = Inches(typo.margin_bottom_inches)
        section.left_margin = Inches(typo.margin_left_inches)
        section.right_margin = Inches(typo.margin_right_inches)
        if doc_title:
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = ""
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = hp.add_run(doc_title[:72])
            _set_run_font(run, size=Pt(8), color=palette.text_muted)
            run.italic = True

    normal = doc.styles["Normal"]
    normal.font.name = typo.body_font
    normal.font.size = Pt(typo.body_size_pt)
    normal.font.color.rgb = palette.text
    nf = normal.paragraph_format
    nf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    nf.line_spacing = typo.line_spacing
    nf.space_after = Pt(typo.space_after_body_pt)
    nf.space_before = Pt(0)
    nf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if _first_line_indent_enabled():
        nf.first_line_indent = Inches(typo.first_line_indent_inches)

    heading_specs = {
        1: (Pt(typo.h1_size_pt), palette.h1_color, Pt(typo.h1_space_before_pt), Pt(typo.h1_space_after_pt), True),
        2: (Pt(typo.h2_size_pt), palette.h2_color, Pt(typo.h2_space_before_pt), Pt(typo.h2_space_after_pt), True),
        3: (Pt(typo.h3_size_pt), palette.h3_color, Pt(typo.h3_space_before_pt), Pt(typo.h3_space_after_pt), True),
    }
    for level, (size, color, before, after, bold) in heading_specs.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = typo.heading_font
        style.font.size = size
        style.font.color.rgb = color
        style.font.bold = bold
        hf = style.paragraph_format
        hf.space_before = before
        hf.space_after = after
        hf.keep_with_next = True

    for list_style in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2"):
        try:
            ls = doc.styles[list_style]
            ls.font.name = typo.body_font
            ls.font.size = Pt(typo.body_size_pt)
            ls.font.color.rgb = palette.text
            ls.paragraph_format.space_after = Pt(3)
        except KeyError:
            continue


def add_cover_header_band(doc: Document, *, height_inches: float = 0.55) -> None:
    """Accent band at the top of the cover page."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, _p().header_band_fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Approximate band height via empty run + spacing
    run = p.add_run(" ")
    run.font.size = Pt(int(height_inches * 28))


def style_cover_title(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.paragraph_format.space_after = Pt(6)
    typo = _typo()
    for run in paragraph.runs:
        _set_run_font(run, name=typo.heading_font, size=Pt(typo.cover_title_size_pt), color=_p().cover_title)
        run.bold = True


def style_cover_subtitle(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    for run in paragraph.runs:
        _set_run_font(run, size=Pt(_typo().cover_subtitle_size_pt), color=_p().text_muted)


def format_metadata_table(table: Table) -> None:
    """Metadata card: soft header row, alternating fills, no harsh grid."""
    _remove_table_borders(table)
    if not table.rows:
        return
    for ci, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, _p().table_header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, size=Pt(10), color=_p().table_header_text)
                run.bold = True
    for ri, row in enumerate(table.rows[1:], start=1):
        for ci, cell in enumerate(row.cells):
            fill = _p().table_alt_fill if ri % 2 else "FFFFFF"
            _set_cell_shading(cell, fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    if ci == 0:
                        _set_run_font(run, size=Pt(10), color=_p().text_muted)
                        run.bold = True
                    else:
                        _set_run_font(run, size=Pt(10), color=_p().text)


def style_toc_title(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    typo = _typo()
    for run in paragraph.runs:
        _set_run_font(run, name=typo.heading_font, size=Pt(typo.toc_title_size_pt), color=_p().toc_title)
        run.bold = True


def _emphasize_heading_runs(
    paragraph,
    *,
    size: Pt,
    color: RGBColor,
    font_name: str | None = None,
) -> None:
    """Force bold + explicit point size on every run (Word heading styles can look faint)."""
    for run in paragraph.runs:
        _set_run_font(run, name=font_name or _heading_font(), size=size, color=color)
        run.bold = True


def style_chapter_heading(paragraph) -> None:
    typo = _typo()
    paragraph.paragraph_format.space_before = Pt(typo.h1_space_before_pt)
    paragraph.paragraph_format.space_after = Pt(typo.h1_space_after_pt)
    _emphasize_heading_runs(paragraph, size=Pt(typo.h1_size_pt), color=_p().h1_color)
    _set_paragraph_border(paragraph, color=_p().h1_border, size=14)


def style_section_heading(paragraph, *, level: int = 2) -> None:
    pal = _p()
    typo = _typo()
    if level == 2:
        color, size, border = pal.h2_color, Pt(typo.h2_size_pt), pal.h2_border
    else:
        color, size, border = pal.h3_color, Pt(typo.h3_size_pt), pal.h3_border
    _emphasize_heading_runs(paragraph, size=size, color=color)
    _set_paragraph_border(paragraph, color=border, size=10)


def style_body_paragraph(paragraph) -> None:
    typo = _typo()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = typo.line_spacing
    paragraph.paragraph_format.space_after = Pt(typo.space_after_body_pt)
    if _first_line_indent_enabled():
        paragraph.paragraph_format.first_line_indent = Inches(typo.first_line_indent_inches)
    for run in paragraph.runs:
        if run.font.size is None:
            _set_run_font(run, size=Pt(typo.body_size_pt), color=_p().text)


def style_bullet_paragraph(paragraph, *, level: int = 0) -> None:
    typo = _typo()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = typo.line_spacing
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Inches(0.2 + 0.12 * level)
    for run in paragraph.runs:
        if run.font.size is None:
            _set_run_font(run, size=Pt(typo.body_size_pt), color=_p().text)


def style_numbered_paragraph(paragraph, *, level: int = 0) -> None:
    style_bullet_paragraph(paragraph, level=level)
    if paragraph.runs:
        paragraph.runs[0].bold = True


def restart_numbered_paragraph(paragraph, *, level: int = 0) -> None:
    """No-op: Word rejects inline ``w:lvlOverride`` on paragraph ``w:numPr``.

    Ordered lists are rendered as plain ``N. text`` paragraphs (numbers come from
    markdown renumbering per section). Kept for API compatibility.
    """
    del paragraph, level  # unused


class NumberedListTracker:
    """Track when the next numbered paragraph at each indent level should restart."""

    def __init__(self) -> None:
        self._pending: Dict[int, bool] = {0: True, 1: True, 2: True}

    def on_section_boundary(self) -> None:
        """New chapter/section — the next list at every level restarts."""
        self._pending = {0: True, 1: True, 2: True}

    def on_non_ordered_line(self) -> None:
        """Prose, bullets, or headings break the current numbered list."""
        self._pending = {0: True, 1: True, 2: True}

    def apply_restart_if_needed(self, paragraph, *, level: int = 0) -> None:
        if self._pending.get(level, True):
            restart_numbered_paragraph(paragraph, level=level)
            self._pending[level] = False


def _callout_style_for_label(label: str) -> tuple[str, str, RGBColor]:
    key = (label or "").strip().lower().rstrip(":")
    pal = _p()
    return pal.callout_styles.get(key, pal.callout_default)


def insert_word_field(paragraph, field_code: str, *, placeholder: str = "") -> None:
    """Insert a Word field (PAGE, NUMPAGES, PAGEREF, etc.) into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    if placeholder:
        text_el = OxmlElement("w:t")
        text_el.text = placeholder
        r.append(text_el)
    r.append(fld_end)


def add_page_number_footer(doc: Document) -> None:
    """Centered 'Page X of Y' footer on every page."""
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Page ")
        insert_word_field(p, "PAGE", placeholder="1")
        p.add_run(" of ")
        insert_word_field(p, "NUMPAGES", placeholder="1")


def enable_update_fields_on_open(doc: Document) -> None:
    """Ask Word to refresh PAGE/NUMPAGES fields when the document opens."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def finalize_word_document(doc: Document) -> None:
    """Apply footer page numbers and auto field refresh."""
    add_page_number_footer(doc)
    enable_update_fields_on_open(doc)


def refresh_word_fields(docx_path: str) -> None:
    """Update PAGE/NUMPAGES fields via Word on Windows."""
    try:
        import win32com.client  # type: ignore[import-untyped]
    except ImportError:
        logger.warning(
            "pywin32 not installed — page numbers may show placeholders until Word "
            "updates fields (Ctrl+A, F9)."
        )
        return

    resolved = str(Path(docx_path).resolve())
    word = None
    doc = None
    wd_print_view = 3
    try:
        word = win32com.client.Dispatch("Word.Application")
        try:
            word.Visible = False
        except AttributeError:
            pass
        word.DisplayAlerts = 0
        doc = word.Documents.Open(resolved)
        doc.ActiveWindow.View.Type = wd_print_view
        doc.Repaginate()
        doc.Fields.Update()
        story = doc.StoryRanges(1)
        while story is not None:
            story.Fields.Update()
            story = story.NextStoryRange
        doc.Save()
        logger.info("Updated Word fields (page numbers): %s", docx_path)
    except Exception as exc:
        logger.warning("Could not auto-update Word fields: %s", exc)
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=True)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def style_toc_entry(paragraph, *, level: int) -> None:
    pal = _p()
    typo = _typo()
    for run in paragraph.runs:
        if level == 1:
            _set_run_font(run, name=typo.body_font, size=Pt(typo.h2_size_pt - 1), color=pal.h1_color)
            run.bold = True
        elif level == 2:
            _set_run_font(run, name=typo.body_font, size=Pt(typo.body_size_pt), color=pal.h2_color)
            run.bold = True
        else:
            _set_run_font(run, name=typo.body_font, size=Pt(typo.body_size_pt - 1), color=pal.h3_color)
            run.bold = True


_NAMED_CALLOUT_LABELS = frozenset(
    {
        "course outcomes",
        "learning objectives",
        "key points",
        "quick revision",
        "definition",
        "important",
        "note",
        "summary",
        "exam tip",
    }
)


def is_callout_label(text: str) -> bool:
    return bool(_CALLOUT_LABEL_RE.match((text or "").strip()))


def is_named_callout_label(label: str) -> bool:
    return (label or "").strip().lower().rstrip(":") in _NAMED_CALLOUT_LABELS


def add_topic_subheading(doc: Document, label: str) -> None:
    """Section topic label (paragraph-style notes) — prominent, not a shaded callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(label.rstrip(":"))
    _set_run_font(run, name=_heading_font(), size=Pt(_typo().h3_size_pt + 1), color=_p().h2_color)
    run.bold = True
    _set_paragraph_border(p, color=_p().h3_border, size=6)


def callout_label_text(text: str) -> str:
    m = _CALLOUT_LABEL_RE.match((text or "").strip())
    return m.group(1).strip() if m else (text or "").strip()


def add_callout_paragraph(doc: Document, text: str, *, label: Optional[str] = None) -> None:
    """Shaded highlight box for key points, course outcomes, definitions, etc."""
    fill, border, label_color = _callout_style_for_label(label or "")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    _set_paragraph_shading(p, fill)
    _set_paragraph_border(p, color=border, size=10, sides=("left",))
    if label:
        lr = p.add_run(f"{label.rstrip(':')} ")
        _set_run_font(lr, size=Pt(11), color=label_color)
        lr.bold = True
    if text:
        tr = p.add_run(text.strip())
        _set_run_font(tr, size=Pt(10), color=_p().text)


def format_content_table(table: Table) -> None:
    """Style markdown tables with a tinted header row."""
    if not table.rows:
        return
    table.style = "Table Grid"
    for ci, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, _p().table_header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, size=Pt(10), color=_p().table_header_text)
                run.bold = True
    for ri, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            if ri % 2 == 0:
                _set_cell_shading(cell, _p().table_alt_fill)
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None:
                        _set_run_font(run, size=Pt(10), color=_p().text)
