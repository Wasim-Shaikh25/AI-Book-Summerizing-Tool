"""Tests for Word document theme styling."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from src.modules.export.docx_notes_exporter import DocxNotesExporter
from src.modules.export.document_formatter import BookCoverMeta
from src.modules.export.docx_theme import (
    add_page_number_footer,
    apply_study_notes_theme,
    finalize_word_document,
    is_callout_label,
    resolve_docx_theme,
)
from src.modules.export.docx_theme_palettes import get_palette
from src.modules.export.markdown_docx_renderer import export_markdown_file_to_docx


def test_apply_study_notes_theme_sets_heading_colors() -> None:
    doc = Document()
    apply_study_notes_theme(doc, doc_title="Sample Book", theme="color")
    h1 = doc.styles["Heading 1"]
    assert h1.font.size.pt == 20
    assert str(h1.font.color.rgb) == "5C2D91"


def test_apply_study_notes_theme_bw() -> None:
    doc = Document()
    apply_study_notes_theme(doc, doc_title="Sample Book", theme="bw")
    h1 = doc.styles["Heading 1"]
    assert str(h1.font.color.rgb) == "000000"
    assert get_palette("bw").name == "bw"


def test_resolve_docx_theme_env(monkeypatch) -> None:
    monkeypatch.setenv("DOCX_THEME", "bw")
    assert resolve_docx_theme() == "bw"
    monkeypatch.setenv("DOCX_THEME", "color")
    assert resolve_docx_theme() == "color"


def test_finalize_word_document_adds_page_footer() -> None:
    doc = Document()
    doc.add_paragraph("Body")
    finalize_word_document(doc)
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "Page" in footer_text


def test_is_callout_label_detects_key_points() -> None:
    assert is_callout_label("**Key Points**")
    assert is_callout_label("**Quick Revision:**")
    assert not is_callout_label("Normal paragraph")


def test_structured_export_applies_theme(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(
        "src.modules.export.docx_notes_exporter.refresh_word_fields",
        lambda _path: None,
    )
    hierarchy = {
        "chapters": [
            {
                "chapter_id": "C1",
                "heading": "Fundamental Rights",
                "sections": [{"section_id": "S1", "heading": "Article 14"}],
            }
        ]
    }
    rewritten = {"S1": "**Key Points**\n\n- Equality before law\n\n> Important case law summary"}
    out = tmp_path / "styled.docx"
    DocxNotesExporter().export(
        out,
        cover=BookCoverMeta(title="Constitution Notes", chapter_count=1, section_count=1),
        hierarchy=hierarchy,
        rewritten=rewritten,
        compact_toc=True,
    )
    assert out.exists()
    doc = Document(str(out))
    assert doc.styles["Normal"].font.name == "Times New Roman"
    assert len(doc.paragraphs) > 5


def test_markdown_export_applies_theme(tmp_path: Path) -> None:
    md = """<div style="text-align: center;">

# Constitution Notes

### Study Notes

</div>

| | |
|:---|:---|
| **Book** | Constitution Notes |
| **Generated** | June 7, 2026 |

# Table of Contents

## 1. Rights

# Fundamental Rights

## Article 14

**Key Points**

- Equality before law

> Courts apply reasonable classification test
"""
    out = tmp_path / "from_md.docx"
    export_markdown_file_to_docx(md, out)
    assert out.exists()
    doc = Document(str(out))
    assert doc.styles["Heading 1"].font.size.pt == 20
