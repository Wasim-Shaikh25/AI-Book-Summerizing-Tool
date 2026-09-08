"""Tests for line-by-line note body DOCX rendering."""

from __future__ import annotations

from docx import Document

from src.modules.export.note_body_docx import append_note_body_markdown


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs if p.text.strip()]


def test_bullets_render_as_separate_list_paragraphs() -> None:
    doc = Document()
    append_note_body_markdown(
        doc,
        "- First point\n- Second point\nPlain line after bullets",
    )
    texts = _paragraph_texts(doc)
    assert "First point" in texts[0]
    assert "Second point" in texts[1]
    assert texts[2] == "Plain line after bullets"
    assert doc.paragraphs[0].style.name.startswith("List Bullet")
    assert doc.paragraphs[1].style.name.startswith("List Bullet")


def test_indented_bullets_still_become_separate_items() -> None:
    """Normalizer flattens nested bullets; each line must still be its own list paragraph."""
    doc = Document()
    append_note_body_markdown(
        doc,
        "- Top\n  - Nested item",
    )
    texts = _paragraph_texts(doc)
    assert texts == ["Top", "Nested item"]
    assert all(p.style.name.startswith("List Bullet") for p in doc.paragraphs[:2])


def test_numbered_sublist_after_bullet() -> None:
    doc = Document()
    append_note_body_markdown(
        doc,
        "- Main idea\n  1. Step one\n  2. Step two",
    )
    assert doc.paragraphs[0].style.name.startswith("List Bullet")
    assert "1. Step one" in doc.paragraphs[1].text
    assert "2. Step two" in doc.paragraphs[2].text


def test_standalone_bold_becomes_subtopic() -> None:
    doc = Document()
    append_note_body_markdown(doc, "**Mahr meaning**\n- Dowry is not mahr")
    assert "Mahr meaning" in doc.paragraphs[0].text
    assert doc.paragraphs[1].style.name.startswith("List Bullet")


def test_numbered_list_restarts_each_section_body() -> None:
    """Topic B must show 1, 2 — numbers come from markdown, not Word list counter."""
    doc = Document()
    append_note_body_markdown(doc, "1. First\n2. Second\n3. Third")
    append_note_body_markdown(doc, "1. New topic first\n2. New topic second")
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0].startswith("1. First")
    assert texts[3].startswith("1. New topic first")
    assert texts[4].startswith("2. New topic second")
    # Must not use Word List Number (invalid restart XML was breaking Word open).
    assert not any(p.style.name.startswith("List Number") for p in doc.paragraphs)


def test_restart_numbered_paragraph_is_noop() -> None:
    from src.modules.export.docx_theme import restart_numbered_paragraph

    doc = Document()
    p = doc.add_paragraph("item", style="List Number")
    restart_numbered_paragraph(p, level=0)
    num_pr = p._p.pPr.numPr if p._p.pPr is not None else None
    if num_pr is not None:
        from docx.oxml.ns import qn

        assert num_pr.find(qn("w:lvlOverride")) is None
